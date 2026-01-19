import streamlit as st
import sqlite3
import uuid
import json
import os
from datetime import datetime
from langchain_openai import ChatOpenAI
from langchain_community.document_loaders import PyPDFLoader
from langchain_text_splitters import RecursiveCharacterTextSplitter
import streamlit as st
from langchain_openai import ChatOpenAI
from langchain_community.document_loaders import PyPDFLoader
from langchain_text_splitters import RecursiveCharacterTextSplitter
from datetime import datetime, timedelta
from reportlab.lib.pagesizes import letter
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
from reportlab.lib.enums import TA_CENTER, TA_JUSTIFY
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
import io
import re
from urllib.parse import quote

st.markdown("""
    <style>
    h1, h2, h3 {
        font-family: 'Utopia', 'Georgia', serif !important;
        font-weight: 600 !important;
    }
    .stApp > header + div [data-testid="stMarkdownContainer"] p {
        font-family: 'Utopia', 'Georgia', serif !important;
    }
    /* Chat messages and general text */
    .stChatMessage, .stMarkdown, p, div {
        font-family: 'Utopia', 'Georgia', serif !important;
    }
    </style>
""", unsafe_allow_html=True)

DB_PATH = "impact_story.db"

def get_conn():
    return sqlite3.connect(DB_PATH, check_same_thread=False)

def init_db():
    conn = get_conn()
    cur = conn.cursor()

    cur.execute("""
    CREATE TABLE IF NOT EXISTS papers (
        paper_id TEXT PRIMARY KEY,
        title TEXT,
        pdf_text TEXT
    )
    """)

    cur.execute("""
    CREATE TABLE IF NOT EXISTS interviews (
        interview_code TEXT PRIMARY KEY,
        paper_id TEXT
    )
    """)

    cur.execute("""
    CREATE TABLE IF NOT EXISTS interview_states (
        interview_code TEXT PRIMARY KEY,
        state_json TEXT,
        transcript_json TEXT
    )
    """)

    conn.commit()

llm = ChatOpenAI(
    model="gpt-4.1-mini",
    api_key=st.secrets["OPENAI_API_KEY"],
    base_url="https://ai-research-proxy.azurewebsites.net"
)

# --------------------------------
# Session state initialization
# --------------------------------

IMPACT_STORY_SECTIONS = {
    "title_hook": "Title",
    "societal_problem": "Introduction",
    "societal_impact": "Societal Impact",
    "research_and_approach": "Research and Approach",
    "people_and_collaboration": "People and Collaboration",
    "outlook": "Conclusion/Outlook"
}

if "paper_text" not in st.session_state:
    st.session_state.paper_text = None

if "impact_state" not in st.session_state:
    st.session_state.impact_state = {
        "title_hook": {"content": None, "answers": [], "attempts": 0},
        "societal_problem": {"content": None, "answers": [], "attempts": 0},
        "societal_impact": {"content": None, "answers": [], "attempts": 0},
        "research_and_approach": {"content": None, "answers": [], "attempts": 0},
        "people_and_collaboration": {"content": None, "answers": [], "attempts": 0},
        "outlook": {"content": None, "answers": [], "attempts": 0}
    }

if "interview_transcript" not in st.session_state:
    st.session_state.interview_transcript = []

if "current_prompt" not in st.session_state:
    st.session_state.current_prompt = None

if "generated_story" not in st.session_state:
    st.session_state.generated_story = None

if "story_revisions" not in st.session_state:
    st.session_state.story_revisions = []  # List of {"story": "...", "revisions": [{"user": "...", "assistant": "..."}]}

if "revision_mode" not in st.session_state:
    st.session_state.revision_mode = False

if "revision_feedback" not in st.session_state:
    st.session_state.revision_feedback = []

if "revision_welcome_shown" not in st.session_state:
    st.session_state.revision_welcome_shown = False

if "admin_authenticated" not in st.session_state:
    st.session_state.admin_authenticated = False
    
MAX_ATTEMPTS_PER_SECTION = 2

EMAIL_SUBJECT = "Research Summary for Impact Story"
EMAIL_INTRO_TEXT = """Dear Colleague,

I am pleased to share the following the summary of the research for impact story. 

Best regards,
Bart van Zelst
Faculty of Economics and Business
Marketing & Communication
"""

STORY_GENERATION_PROMPT = f"""Write a complete Impact Story for a broad audience using the structure below.

Structure and word limits:
- {IMPACT_STORY_SECTIONS["title_hook"]} (short, active)
- {IMPACT_STORY_SECTIONS["societal_problem"]} (~100 words)
- {IMPACT_STORY_SECTIONS["societal_impact"]} (~150 words)
- {IMPACT_STORY_SECTIONS["research_and_approach"]} (~150 words)
- {IMPACT_STORY_SECTIONS["people_and_collaboration"]} (~100 words)
- {IMPACT_STORY_SECTIONS["outlook"]} (~50 words)

Keep the story concrete and human. Avoid jargon.
Focus on the connection between science and everyday life.

Use the following elicited content as your primary source:

"""

def summarize_paper_sections(full_text):
    prompt = (
        "You are an academic assistant.\n\n"
        "Summarize the following research paper into the following sections:\n"
        "1. Research problem and motivation\n"
        "2. Methodology and data\n"
        "3. Key findings\n"
        "4. Theoretical or scientific contribution\n"
        "5. Practical or real-world implications\n\n"
        "Paper text:\n"
        f"{full_text}"
    )
    return llm.invoke(prompt).content

def next_missing_section(state):
    for k, v in state.items():
        if v["content"] is None and v["attempts"] < MAX_ATTEMPTS_PER_SECTION:
            return k
    return None

def get_last_researcher_answer(transcript):
    for msg in reversed(transcript):
        if msg["role"] == "researcher":
            return msg["content"]
    return None

def generate_interview_turn(section, paper_text, last_answer=None):
    prompt = (
        "You are facilitating an impact story interview for a broad, non-academic audience.\n\n"
        "First, briefly acknowledge and paraphrase the researcher's previous answer "
        "in plain language (if one exists).\n"
        "Then explain why this aspect matters from a societal or human perspective.\n"
        "Finally, ask ONE open-ended question to help develop the following impact story section:\n\n"
        f"{section.replace('_', ' ').title()}\n\n"
        "Avoid jargon. Keep the tone conversational and supportive.\n\n"
        "Research paper context:\n"
        f"{paper_text}\n\n"
        f"Previous answer:\n{last_answer if last_answer else 'N/A'}"
    )
    return llm.invoke(prompt).content

def evaluate_and_store(section, state, paper_text):
    section_data = state[section]
    combined_answers = "\n".join(section_data["answers"])

    prompt = (
        f"You are helping synthesize an impact story section: "
        f"{section.replace('_', ' ').title()}.\n\n"
        "Below are multiple responses provided by the researcher over the interview.\n\n"
        f"{combined_answers}\n\n"
        "If this information is sufficient, synthesize a clear, human-friendly "
        "paragraph (3–4 sentences) suitable for a broad audience.\n"
        "If information is still incomplete, return ONLY the word 'INSUFFICIENT'."
    )

    result = llm.invoke(prompt).content.strip()
    
     # If sufficient, use it
    if result.upper() != "INSUFFICIENT":
        section_data["content"] = result
        return result

    # If insufficient AND we've hit max attempts, force a best-effort fallback
    if section_data["attempts"] >= MAX_ATTEMPTS_PER_SECTION:
        fallback_prompt = (
            f"Based on the partial information below, write the best possible "
            f"version of the impact story section '{section.replace('_',' ')}'. "
            "Be transparent about what is unknown, but still produce a usable paragraph.\n\n"
            f"{combined_answers}"
        )
        fallback = llm.invoke(fallback_prompt).content.strip()
        section_data["content"] = fallback
        return fallback

    # Otherwise, keep asking
    return None

def extract_title(story_text):
    """Extract and sanitize the title from the story text for use in filenames."""
    paragraphs = story_text.split('\n\n')
    if paragraphs:
        title = paragraphs[0].strip()
        # Remove markdown formatting
        title = re.sub(r'\*\*(.+?)\*\*', r'\1', title)
        title = re.sub(r'\*([^*]+?)\*', r'\1', title)
        # Sanitize for filename (remove invalid characters)
        title = re.sub(r'[<>:"/\\|?*]', '', title)
        # Limit length and clean up
        title = title[:100].strip()
        return title if title else "impact_story"
    return "impact_story"


def generate_pdf(story_text):
    """Generate a PDF document from the impact story text."""
    buffer = io.BytesIO()
    doc = SimpleDocTemplate(buffer, pagesize=letter,
                            rightMargin=72, leftMargin=72,
                            topMargin=72, bottomMargin=18)
    
    styles = getSampleStyleSheet()
    story = []
    
    # Title style
    title_style = ParagraphStyle(
        'CustomTitle',
        parent=styles['Heading1'],
        fontSize=18,
        textColor='#1f77b4',
        spaceAfter=30,
        alignment=TA_CENTER
    )
    
    # Normal style with justified text
    normal_style = ParagraphStyle(
        'CustomNormal',
        parent=styles['Normal'],
        fontSize=11,
        leading=14,
        alignment=TA_JUSTIFY,
        spaceAfter=12
    )
    
    # Split story into paragraphs
    paragraphs = story_text.split('\n\n')
    
    for i, para in enumerate(paragraphs):
        para = para.strip()
        if not para:
            continue
        
        # Check if it's a title (usually first paragraph or all caps/short)
        is_title = (len(para) < 100 and para.isupper()) or (i == 0)
        
        if is_title:
            # Remove markdown formatting from title
            title_clean = re.sub(r'\*\*(.+?)\*\*', r'\1', para)
            title_clean = re.sub(r'\*([^*]+?)\*', r'\1', title_clean)
            story.append(Paragraph(title_clean, title_style))
        else:
            # Replace markdown formatting for PDF (ReportLab uses HTML-like tags)
            # Handle bold: **text** -> <b>text</b>
            para = re.sub(r'\*\*(.+?)\*\*', r'<b>\1</b>', para)
            # Handle italic: *text* -> <i>text</i> (but not if it's part of **)
            para = re.sub(r'(?<!\*)\*([^*]+?)\*(?!\*)', r'<i>\1</i>', para)
            story.append(Paragraph(para, normal_style))
        
        story.append(Spacer(1, 0.2*inch))
    
    doc.build(story)
    buffer.seek(0)
    return buffer.getvalue()


def generate_word(story_text):
    """Generate a Word document from the impact story text."""
    doc = Document()
    
    # Set default font
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)
    
    # Split story into paragraphs
    paragraphs = story_text.split('\n\n')
    
    for i, para in enumerate(paragraphs):
        para = para.strip()
        if not para:
            continue
        
        # Check if it's a title
        is_title = (len(para) < 100 and para.isupper()) or (i == 0)
        
        if is_title:
            # Remove markdown formatting from title
            title_clean = re.sub(r'\*\*(.+?)\*\*', r'\1', para)
            title_clean = re.sub(r'\*([^*]+?)\*', r'\1', title_clean)
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(title_clean)
            run.font.size = Pt(18)
            run.font.bold = True
        else:
            # Handle markdown formatting for Word
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            
            # Process text with markdown formatting
            text_parts = re.split(r'(\*\*.*?\*\*|\*.*?\*)', para)
            for part in text_parts:
                if not part:
                    continue
                if part.startswith('**') and part.endswith('**'):
                    # Bold text
                    run = p.add_run(part[2:-2])
                    run.bold = True
                elif part.startswith('*') and part.endswith('*') and not part.startswith('**'):
                    # Italic text
                    run = p.add_run(part[1:-1])
                    run.italic = True
                else:
                    # Regular text
                    p.add_run(part)
        
        doc.add_paragraph()  # Add spacing between paragraphs
    
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer.getvalue()


def generate_email_content(story_text, subject, intro_text):
    """Generate email content with subject, intro text, and story."""
    # Remove markdown formatting from story for email (keep it simple)
    story_clean = re.sub(r'\*\*(.+?)\*\*', r'\1', story_text)
    story_clean = re.sub(r'\*([^*]+?)\*', r'\1', story_clean)
    
    email_body = f"{intro_text}\n\n{story_clean}"
    return subject, email_body


def create_mailto_link(subject, body):
    """Create a mailto link with subject and body."""
    subject_encoded = quote(subject)
    body_encoded = quote(body)
    return f"mailto:?subject={subject_encoded}&body={body_encoded}"

def compute_interview_status(state_json):
    state = json.loads(state_json)

    completed = sum(
        1 for v in state.values() if v.get("content") is not None
    )
    total = len(state)

    if completed == 0:
        return "Not started"
    elif completed < total:
        return f"In progress ({completed}/{total})"
    else:
        return "Completed"
        
def admin_panel():
    st.header("Admin Panel")

    if not st.session_state.admin_authenticated:
        pwd = st.text_input(
            "Enter admin password",
            type="password"
        )

        if st.button("Login"):
            if pwd == st.secrets["ADMIN_PASSWORD"]:
                st.session_state.admin_authenticated = True
                st.success("Access granted")
                st.rerun()
            else:
                st.error("Invalid password")

        return  # stop here if not authenticated
    if st.session_state.admin_authenticated:

        if st.button("Logout"):
            st.session_state.admin_authenticated = False
            st.rerun()
        
        uploaded_file = st.file_uploader("Upload research paper (PDF)", type=["pdf"])
    
        if uploaded_file and st.button("Process paper"):
            with st.spinner("Reading and processing PDF..."):
                paper_id = str(uuid.uuid4())
                interview_code = str(uuid.uuid4())[:8]
                
                with open("temp.pdf", "wb") as f:
                    f.write(uploaded_file.getbuffer())
        
                loader = PyPDFLoader("temp.pdf")
                documents = loader.load()
        
                splitter = RecursiveCharacterTextSplitter(
                    chunk_size=1500,
                    chunk_overlap=200
                )
                docs = splitter.split_documents(documents)
        
                paper_text = "\n\n".join(d.page_content for d in docs)
                text = summarize_paper_sections(paper_text)
    
                conn = get_conn()
                conn.execute(
                    "INSERT INTO papers VALUES (?, ?, ?)",
                    (paper_id, uploaded_file.name, text)
                )
        
                conn.execute(
                    "INSERT INTO interviews VALUES (?, ?)",
                    (interview_code, paper_id)
                )
        
                conn.execute(
                    "INSERT INTO interview_states VALUES (?, ?, ?)",
                    (
                        interview_code,
                        json.dumps(st.session_state.impact_state),
                        json.dumps([])
                    )
                )
        
                conn.commit()
                conn.close()
        
                st.success("Interview created")
                st.code(f"Researcher access code: {interview_code}")
    
        conn = get_conn()
    
        rows = conn.execute("""
            SELECT 
                i.interview_code,
                p.title,
                s.state_json
            FROM interviews i
            JOIN papers p ON i.paper_id = p.paper_id
            JOIN interview_states s ON i.interview_code = s.interview_code
        """).fetchall()
        
        conn.close()
    
        st.subheader("📋 All Interviews")
    
        table_data = []
        
        for interview_code, title, state_json in rows:
            status = compute_interview_status(state_json)
        
            table_data.append({
                "Paper": title,
                "Interview Code": interview_code,
                "Status": status
            })
        
        st.table(table_data)

def researcher_login():
    st.header("Researcher Login")
    code = st.text_input("Enter access code")

    if st.button("Access interview"):
        conn = get_conn()
        row = conn.execute(
            "SELECT interview_code FROM interviews WHERE interview_code = ?",
            (code,)
        ).fetchone()
        conn.close()

        if row:
            st.session_state.interview_code = code
            st.rerun()
        else:
            st.error("Invalid code")

# ---------------- INTERVIEW ----------------
def interview_ui():

    code = st.session_state.interview_code

    conn = get_conn()
    paper_text = conn.execute("""
        SELECT p.pdf_text
        FROM papers p
        JOIN interviews i ON p.paper_id = i.paper_id
        WHERE i.interview_code = ?
    """, (code,)).fetchone()[0]

    row = conn.execute(
            "SELECT state_json, transcript_json FROM interview_states WHERE interview_code = ?",
            (code,)
        ).fetchone()

    state = json.loads(row[0])
    st.session_state.impact_state = state  # sync session with DB state

    if row[1]:
        st.session_state.interview_transcript = json.loads(row[1])
    else:
        st.session_state.interview_transcript = []
    
    conn.close()

    with st.sidebar.expander("Question Progress", expanded=True):
        impact_state = state
        section_keys = list(IMPACT_STORY_SECTIONS.keys())
        
        for idx, (section_key, display_name) in enumerate(IMPACT_STORY_SECTIONS.items(), 1):
            section_data = impact_state.get(section_key, {})
            is_answered = section_data.get("content") is not None
            
            if is_answered:
                st.markdown(f"✅ {display_name}")
                st.caption("Answered")
            else:
                # Find current section being worked on
                missing_section = next(
                    (k for k, v in impact_state.items() 
                     if v.get("content") is None and v.get("attempts", 0) < MAX_ATTEMPTS_PER_SECTION),
                    None
                )
                
                if section_key == missing_section:
                    st.markdown(f"🔴 {display_name}")
                    st.caption(f"In progress")
                else:
                    # Find the position of this section among unanswered sections
                    unanswered_sections = [k for k in IMPACT_STORY_SECTIONS.keys() 
                                          if impact_state.get(k, {}).get("content") is None]
                    position = unanswered_sections.index(section_key) + 1 if section_key in unanswered_sections else idx
                    st.markdown(f"⚪ {display_name}")
                    st.caption(f"Pending")
                    
    missing_section = next_missing_section(state)

    all_answered = all(v["content"] is not None for v in state.values())
                
    if missing_section:
        if st.session_state.current_prompt is None:
            last_answer = get_last_researcher_answer(
                st.session_state.interview_transcript
            )

            with st.spinner("Preparing next interview turn..."):
                st.session_state.current_prompt = generate_interview_turn(
                    missing_section,
                    paper_text,
                    last_answer
                )
                
    for msg in st.session_state.interview_transcript:
        role = "assistant" if msg["role"] == "assistant" else "user"
    
        with st.chat_message(role):
            st.write(msg["content"])

    if missing_section and st.session_state.current_prompt:
        with st.chat_message("assistant"):
            st.write(st.session_state.current_prompt)

    st.markdown("""
        <style>
        /* Hide the default send icon and show "Send" text instead */
        div[data-testid="stChatInput"] button {
            position: relative;
            min-width: 80px !important;
            min-height: 40px !important;
            padding: 8px 16px !important;
            background-color: #bc0031 !important;
            border-color: #bc0031 !important;
        }
        div[data-testid="stChatInput"] button:hover {
            background-color: #d1003a !important;
            border-color: #d1003a !important;
        }
        div[data-testid="stChatInput"] button svg {
            display: none !important;
        }
        div[data-testid="stChatInput"] button::after {
            content: "Send";
            display: inline-block;
            font-size: 14px;
            font-weight: 500;
            color: white !important;
        }
        /* Style all buttons except chat input button with black background */
        button[kind="primary"]:not(div[data-testid="stChatInput"] button),
        button[kind="secondary"]:not(div[data-testid="stChatInput"] button),
        .stDownloadButton > button,
        a[data-testid="stLinkButton"],
        .stLinkButton > a {
            background-color: #000000 !important;
            color: #ffffff !important;
            border: 1px solid #000000 !important;
        }
        button[kind="primary"]:not(div[data-testid="stChatInput"] button):hover,
        button[kind="secondary"]:not(div[data-testid="stChatInput"] button):hover,
        .stDownloadButton > button:hover,
        a[data-testid="stLinkButton"]:hover,
        .stLinkButton > a:hover {
            background-color: #333333 !important;
            border-color: #333333 !important;
            color: #ffffff !important;
        }
        </style>
    """, unsafe_allow_html=True)

    if not st.session_state.get("revision_mode", False):
        user_answer = st.chat_input("Type your answer here...", disabled=all_answered)
    else:
        user_answer = None

    if user_answer and missing_section:
        st.session_state.interview_transcript.append({
                "role": "assistant",
                "content": st.session_state.current_prompt,
                "timestamp": datetime.utcnow().isoformat()
            })
            
        st.session_state.interview_transcript.append({
                "role": "researcher",
                "content": user_answer,
                "timestamp": datetime.utcnow().isoformat()
            })


        with st.spinner("Integrating response..."):
            section_data = state[missing_section]
            section_data["answers"].append(user_answer)
            section_data["attempts"] += 1
                
            section_data["content"] = evaluate_and_store(missing_section, state, paper_text)
            st.session_state.impact_state = state  # keep session updated
            
            conn = get_conn()
            conn.execute("""
                    UPDATE interview_states
                    SET state_json = ?, transcript_json = ?
                    WHERE interview_code = ?
                """, (
                    json.dumps(state),
                    json.dumps(st.session_state.interview_transcript),
                    code
                ))
            conn.commit()
            conn.close()
            
        st.session_state.current_prompt = None
        st.rerun()
    
    if all_answered:
        st.success("All impact story elements have been collected.")

        if st.button("Generate Impact Story Summary"):
            with st.spinner("Writing impact story summary..."):
                prompt = STORY_GENERATION_PROMPT

                for k, v in state.items():
                    section_name = IMPACT_STORY_SECTIONS.get(k, k.replace('_', ' ').title())
                    prompt += f"\n{section_name}:\n{v['content']}\n"

                impact_story = llm.invoke(prompt).content
                st.session_state.generated_story = impact_story
                # Store the story
                st.session_state.story_revisions.append({
                    "story": impact_story,
                    "revisions": []
                })

        # Display all stories and revisions in chronological order
        for item in st.session_state.story_revisions:
            # Display story
            st.markdown(item["story"])
            
            # Display revisions after the story
            for rev in item["revisions"]:
                with st.chat_message("user"):
                    st.write(rev["user"])
                with st.chat_message("assistant"):
                    st.write(rev["assistant"])
        
        # Display current revision interface
        if st.session_state.revision_mode:
            # Show welcome message once
            if not st.session_state.revision_welcome_shown:
                with st.chat_message("assistant"):
                    st.write("Welcome back, please tell me what you would like to change in the summary.")
                st.session_state.revision_welcome_shown = True
            
            # Display current revision conversation
            for feedback in st.session_state.revision_feedback:
                with st.chat_message("user"):
                    st.write(feedback["user_input"])
                with st.chat_message("assistant"):
                    st.write("Thank you for your answer. I have all the necessary information I need and your revision will be incorporated into final summary. Do you want to add anything else or should we end this interview?")
            
            # Chat input for revision feedback
            revision_input = st.chat_input("Type your revision request here...")
            
            if revision_input:
                st.session_state.revision_feedback.append({"user_input": revision_input})
                st.rerun()
        
        # Display action buttons if story is generated
        if st.session_state.generated_story:
            # Extract title for filename
            story_title = extract_title(st.session_state.generated_story)
            
            # Create columns for side-by-side buttons
            col1, col2, col3, col4 = st.columns(4)
            
            with col1:
                if st.session_state.revision_mode:
                    if st.button("End interview and generate summary"):
                        # Regenerate story with revision feedback
                        with st.spinner("Regenerating story with your revisions..."):
                            feedback_text = "\n".join([f["user_input"] for f in st.session_state.revision_feedback])
                            prompt = STORY_GENERATION_PROMPT

                            for k, v in st.session_state.impact_state.items():
                                section_name = IMPACT_STORY_SECTIONS.get(k, k.replace('_', ' ').title())
                                prompt += f"\n{section_name}:\n{v['content']}\n"
                            
                            prompt += f"\n\nRevision requests from the user:\n{feedback_text}\n\n"
                            prompt += "Please incorporate these revision requests into the story."
                            
                            impact_story = llm.invoke(prompt).content
                            st.session_state.generated_story = impact_story
                            
                            # Store revisions with the last story, then add new story
                            if st.session_state.story_revisions:
                                for feedback in st.session_state.revision_feedback:
                                    st.session_state.story_revisions[-1]["revisions"].append({
                                        "user": feedback["user_input"],
                                        "assistant": "Thank you for your answer. I have all the necessary information I need. Do you want to add anything else or should we end this interview?"
                                    })
                            
                            # Add new story
                            st.session_state.story_revisions.append({
                                "story": impact_story,
                                "revisions": []
                            })
                            
                            st.session_state.revision_mode = False
                            st.session_state.revision_welcome_shown = False
                            st.session_state.revision_feedback = []
                            st.rerun()
                else:
                    if st.button("Revise Story"):
                        st.session_state.revision_mode = True
                        st.session_state.revision_welcome_shown = False
                        st.session_state.revision_feedback = []
                        st.rerun()
            
            with col2:
                pdf_bytes = generate_pdf(st.session_state.generated_story)
                st.download_button(
                    label="Download as PDF",
                    data=pdf_bytes,
                    file_name=f"{story_title}.pdf",
                    mime="application/pdf"
                )
            
            with col3:
                word_bytes = generate_word(st.session_state.generated_story)
                st.download_button(
                    label="Download as Word",
                    data=word_bytes,
                    file_name=f"{story_title}.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                )
            
            with col4:
                # Email export
                email_subject, email_body = generate_email_content(
                    st.session_state.generated_story,
                    EMAIL_SUBJECT,
                    EMAIL_INTRO_TEXT
                )
                mailto_link = create_mailto_link(email_subject, email_body)
                st.link_button(
                    "Send via Email",
                    mailto_link
                )

st.set_page_config(
    page_title="IxA",
    page_icon="resources/uva_icon.png",
    layout="wide"
)
st.image("resources/compacte-logo.jpg", width=200)
st.title("Research Impact Interview")
st.caption(
    "A guided human–AI process for translating research into summary to assist the creation of societal impact stories."
)
        
# ---------------- MAIN ----------------
st.set_page_config(page_title="Impact Story Interview", layout="centered")
init_db()

st.sidebar.title("Role")
role = st.sidebar.radio("Select role", ["Researcher", "Admin"])

if role == "Admin":
    admin_panel()
else:
    if "interview_code" not in st.session_state:
        researcher_login()
    else:
        interview_ui()